"""
Generate Word (.docx) templates for VisaBud Premium tier.
Each template includes:
- Locked sections (gov guidance + examples)
- Comments (WHY each section matters)
- Unlocked sections (user editable)
- Professional formatting
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Create templates directory
TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), '..', 'public', 'templates')
os.makedirs(TEMPLATES_DIR, exist_ok=True)


def add_locked_section(doc, title, content, guidance_comment=None):
    """Add a locked (protected) section with optional comment."""
    # Title (locked, bold, blue)
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(12)
    title_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    
    # Content (locked)
    content_para = doc.add_paragraph(content)
    content_para.paragraph_format.left_indent = Inches(0.3)
    
    # Guidance comment (if provided)
    if guidance_comment:
        comment_para = doc.add_paragraph()
        comment_run = comment_para.add_run(f"💡 {guidance_comment}")
        comment_run.font.size = Pt(10)
        comment_run.font.italic = True
        comment_run.font.color.rgb = RGBColor(100, 100, 100)
        comment_para.paragraph_format.left_indent = Inches(0.5)
    
    # Add spacing
    doc.add_paragraph()


def add_editable_section(doc, title, placeholder, guidance=None):
    """Add an unlocked (editable) section for user input."""
    # Title (locked, bold, green)
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(12)
    title_run.font.color.rgb = RGBColor(0, 102, 51)  # Dark green
    
    # Guidance (if provided)
    if guidance:
        guide_para = doc.add_paragraph(guidance)
        guide_para.paragraph_format.left_indent = Inches(0.3)
        guide_run = guide_para.runs[0]
        guide_run.font.size = Pt(10)
        guide_run.font.italic = True
    
    # Placeholder (light gray background to show editable area)
    placeholder_para = doc.add_paragraph(placeholder)
    placeholder_para.paragraph_format.left_indent = Inches(0.3)
    placeholder_para.paragraph_format.space_after = Pt(12)
    
    # Add spacing
    doc.add_paragraph()


def create_spouse_proof_of_relationship():
    """Template 1: Spouse Visa - Proof of Relationship"""
    doc = Document()
    
    # Header
    header = doc.add_paragraph()
    header_run = header.add_run('SPOUSE VISA - PROOF OF RELATIONSHIP')
    header_run.bold = True
    header_run.font.size = Pt(16)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Version info
    version_para = doc.add_paragraph('Version: Premium Template | Last Updated: April 2026')
    version_para.paragraph_format.space_after = Pt(24)
    version_run = version_para.runs[0]
    version_run.font.size = Pt(9)
    version_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Locked: Gov Requirements
    add_locked_section(
        doc,
        '📋 GOVERNMENT REQUIREMENTS',
        (
            'Home Office requires evidence that you have been in a genuine and '
            'subsisting relationship with your partner. Evidence must cover the '
            'entire period of your relationship (or at least the last 2+ years).\n\n'
            'Source: gov.uk/guidance/prove-your-relationship'
        ),
        'Immigration officers check for: consistency across all documents, '
        'timeline gaps, evidence of genuine cohabitation.'
    )
    
    # Locked: Example
    add_locked_section(
        doc,
        '✓ EXAMPLE: What Good Looks Like',
        (
            'My name is John Smith and I have been in a relationship with '
            'Sarah Johnson since January 2021 (5 years). We met at University and '
            'began living together in September 2021. We purchased our home together '
            'in March 2022 and are now planning our wedding for June 2024.\n\n'
            'Evidence supporting this relationship is included as:\n'
            '• Joint tenancy agreement (2021-2022)\n'
            '• Mortgage document with both names (2022-present)\n'
            '• Joint utility bills (2022-present)\n'
            '• Photos together (2021-2024)\n'
            '• Email correspondence (2021-2024)'
        ),
        'Notice the specificity: exact dates, names, timeline, and what evidence '
        'is attached. This is what immigration officers expect.'
    )
    
    # Editable: User section
    add_editable_section(
        doc,
        '✎ YOUR PROOF OF RELATIONSHIP STATEMENT',
        (
            'My name is [YOUR FULL NAME] and I have been in a relationship with '
            '[PARTNER\'S NAME] since [START DATE] ([NUMBER] years/months). '
            'We met [HOW YOU MET] and began living together in [DATE]. '
            '[ADD ANY MAJOR MILESTONES: married, children, property purchase, etc.]\n\n'
            'Evidence supporting this relationship is included as:\n'
            '• [EVIDENCE TYPE 1]\n'
            '• [EVIDENCE TYPE 2]\n'
            '• [EVIDENCE TYPE 3]'
        ),
        'Edit this section with YOUR details. Keep the same structure and style '
        'as the example. Include: names, exact dates, timeline, evidence list.'
    )
    
    # Locked: Why it matters
    add_locked_section(
        doc,
        '⚠️ WHY THIS SECTION MATTERS',
        (
            'Immigration officers must determine if your relationship is "genuine '
            'and subsisting". They check for:\n\n'
            '✓ Dates match across all documents (bank statements, tenancy, photos)\n'
            '✓ Timeline is logical (not meeting and cohabiting same month)\n'
            '✓ Evidence of daily life together (joint bills, photos, correspondence)\n'
            '✓ Consistency with any previous applications/visas\n\n'
            'If dates don\'t match or timeline has unexplained gaps, the '
            'application may be rejected or delayed for further enquiries.'
        ),
        None
    )
    
    # Save
    filename = os.path.join(TEMPLATES_DIR, 'Spouse_ProofOfRelationship.docx')
    doc.save(filename)
    print(f'✅ Created: {filename}')
    return filename


def create_spouse_financial_evidence():
    """Template 2: Spouse Visa - Financial Evidence Summary"""
    doc = Document()
    
    # Header
    header = doc.add_paragraph()
    header_run = header.add_run('SPOUSE VISA - FINANCIAL EVIDENCE SUMMARY')
    header_run.bold = True
    header_run.font.size = Pt(16)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Version info
    version_para = doc.add_paragraph('Version: Premium Template | Last Updated: April 2026')
    version_para.paragraph_format.space_after = Pt(24)
    version_run = version_para.runs[0]
    version_run.font.size = Pt(9)
    version_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Locked: Gov Requirements
    add_locked_section(
        doc,
        '📋 GOVERNMENT FINANCIAL REQUIREMENTS (2024)',
        (
            'Spouse visa applicants must prove their UK sponsor can maintain '
            'themselves and their family without public funds.\n\n'
            'Minimum annual income: £29,000 (2024)\n'
            'Plus additional amount per dependent (£3,800 per child)\n\n'
            'Source: gov.uk/guidance/prove-your-income'
        ),
        'Immigration officers verify: income matches tax returns, P60s, and '
        'payslips; savings meet thresholds if income is lower.'
    )
    
    # Locked: Example
    add_locked_section(
        doc,
        '✓ EXAMPLE: What Good Looks Like',
        (
            'FINANCIAL EVIDENCE SUMMARY\n'
            'Sponsor: John Smith | Dependents: 1 (spouse Sarah Johnson)\n\n'
            'INCOME:\n'
            'Employment: £35,000 annual salary (confirmed by P60 and 3 months payslips)\n'
            'Self-employment: £5,000 annual profit (confirmed by accountant reference)\n'
            'TOTAL ANNUAL INCOME: £40,000\n\n'
            'THRESHOLD REQUIRED: £29,000 (base) + £3,800 (1 dependent) = £32,800\n'
            'STATUS: ✓ Exceeds requirement by £7,200\n\n'
            'SUPPORTING EVIDENCE ATTACHED:\n'
            '• P60 for 2023 (confirming £35,000 salary)\n'
            '• Payslips for Jan, Feb, Mar 2024 (showing ongoing employment)\n'
            '• Accountant reference letter (confirming self-employment income)\n'
            '• Bank statements (6 months, showing income deposits)'
        ),
        'Notice: Clear calculation, evidence list, and status confirmation. '
        'Immigration officers can instantly see if you meet requirements.'
    )
    
    # Editable: User section
    add_editable_section(
        doc,
        '✎ YOUR FINANCIAL EVIDENCE',
        (
            'SPONSOR: [YOUR NAME] | DEPENDENTS: [NUMBER] (spouse [NAME])\n\n'
            'INCOME:\n'
            'Employment: £[AMOUNT] annual salary (source: P60 and recent payslips)\n'
            'Self-employment: £[AMOUNT] annual profit (source: accountant reference)\n'
            '[OTHER INCOME]: £[AMOUNT] (source: [DOCUMENT TYPE])\n'
            'TOTAL ANNUAL INCOME: £[TOTAL]\n\n'
            'THRESHOLD REQUIRED: £29,000 + £[3,800 × NUMBER OF DEPENDENTS] = £[TOTAL REQUIRED]\n'
            'STATUS: ✓ [Exceeds/Meets] requirement by £[DIFFERENCE]\n\n'
            'SUPPORTING EVIDENCE ATTACHED:\n'
            '• [DOCUMENT TYPE 1]\n'
            '• [DOCUMENT TYPE 2]\n'
            '• [DOCUMENT TYPE 3]'
        ),
        'Copy the example format. Replace [BRACKETS] with YOUR actual figures. '
        'Always calculate and show: income vs. requirement vs. difference.'
    )
    
    # Locked: Why it matters
    add_locked_section(
        doc,
        '⚠️ WHY THIS SECTION MATTERS',
        (
            'Financial evidence is one of the top rejection reasons for spouse visas.\n\n'
            'Common mistakes:\n'
            '✗ Not calculating total income correctly (missing self-employment income)\n'
            '✗ Not showing how calculation meets threshold\n'
            '✗ Missing supporting documents (no payslips, no P60)\n'
            '✗ Income figures don\'t match bank statements or tax returns\n\n'
            'If your income doesn\'t meet the threshold:\n'
            '• You can use savings (£16,000 per £1 of shortfall)\n'
            '• You can combine sponsor + partner income\n'
            '• But you MUST document this explicitly'
        ),
        None
    )
    
    # Save
    filename = os.path.join(TEMPLATES_DIR, 'Spouse_FinancialEvidence.docx')
    doc.save(filename)
    print(f'✅ Created: {filename}')
    return filename


def create_skilled_worker_maintenance():
    """Template 3: Skilled Worker - Maintenance of Funds Declaration"""
    doc = Document()
    
    # Header
    header = doc.add_paragraph()
    header_run = header.add_run('SKILLED WORKER VISA - MAINTENANCE OF FUNDS DECLARATION')
    header_run.bold = True
    header_run.font.size = Pt(16)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Version info
    version_para = doc.add_paragraph('Version: Premium Template | Last Updated: April 2026')
    version_para.paragraph_format.space_after = Pt(24)
    version_run = version_para.runs[0]
    version_run.font.size = Pt(9)
    version_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Locked: Gov Requirements
    add_locked_section(
        doc,
        '📋 GOVERNMENT REQUIREMENTS',
        (
            'Skilled Worker visa applicants must prove they have sufficient '
            'funds to maintain themselves in the UK without public funds.\n\n'
            'Minimum amount: £3,100 in savings OR\n'
            'Sponsor can provide maintenance confirmation OR\n'
            'Salary covers your living costs\n\n'
            'Source: gov.uk/guidance/skilled-worker-visa'
        ),
        'Immigration officers check: bank statements match savings claim; '
        'funds are accessible and in your name or joint account.'
    )
    
    # Locked: Example
    add_locked_section(
        doc,
        '✓ EXAMPLE: What Good Looks Like',
        (
            'I, James Wong, am applying for a Skilled Worker visa. I confirm that:\n\n'
            'MAINTENANCE OF FUNDS:\n'
            'I have £5,000 in my UK savings account (NatWest Savings Account, '
            'sort code 60-40-24, account ending 1234).\n'
            'This account has been held for 28 days continuously (as of 14 Apr 2024).\n'
            'Bank statement attached confirms this balance.\n\n'
            'OR: My employer, Acme Tech Ltd, has confirmed in writing that my '
            'salary of £32,500 per annum is sufficient to cover my living costs '
            'in London.\n\n'
            'DECLARATION:\n'
            'I declare that this information is true and complete. I understand '
            'that providing false information may result in visa refusal.\n\n'
            'Signed: James Wong | Date: 14 April 2024'
        ),
        'Clear declaration, specific figures, account details (not full number!), '
        'and attached evidence. Immigration officers can easily verify.'
    )
    
    # Editable: User section
    add_editable_section(
        doc,
        '✎ YOUR MAINTENANCE OF FUNDS DECLARATION',
        (
            'I, [YOUR FULL NAME], am applying for a Skilled Worker visa. I confirm that:\n\n'
            'MAINTENANCE OF FUNDS:\n'
            'I have £[AMOUNT] in my savings account ([BANK NAME] [ACCOUNT TYPE], '
            'sort code [SORT CODE], account ending [LAST 4 DIGITS]).\n'
            'This account has been held for 28 days continuously (as of [DATE]).\n'
            'Bank statement attached confirms this balance.\n\n'
            'OR: My employer, [EMPLOYER NAME], has confirmed in writing that my '
            'salary of £[ANNUAL SALARY] per annum is sufficient to cover my '
            'living costs in [CITY].\n\n'
            'DECLARATION:\n'
            'I declare that this information is true and complete. I understand '
            'that providing false information may result in visa refusal.\n\n'
            'Signed: [YOUR SIGNATURE] | Date: [DATE]'
        ),
        'Edit with YOUR details. IMPORTANT: Do NOT include your full account '
        'number (use last 4 digits only). Ensure 28-day fund rule is met.'
    )
    
    # Locked: Why it matters
    add_locked_section(
        doc,
        '⚠️ WHY THIS SECTION MATTERS',
        (
            'Maintenance of funds rejections are common because applicants:\n\n'
            '✗ Don\'t meet the 28-day continuous fund rule (funds must be held for 28 days)\n'
            '✗ Don\'t provide bank statements showing the funds\n'
            '✗ List funds that aren\'t accessible (tied up in investments, locked savings)\n'
            '✗ Don\'t have evidence they can earn enough to support themselves\n\n'
            'The Home Office will:\n'
            '• Check your bank statement balance matches your claim\n'
            '• Verify the funds were held 28+ days before application\n'
            '• Check that funds are in YOUR name or a joint account (not borrowed)\n\n'
            'If you DON\'T have £3,100 in savings:\n'
            '• Your salary must clearly cover living costs (£2,000+ per month in London)\n'
            '• Your employer must provide a written confirmation letter'
        ),
        None
    )
    
    # Save
    filename = os.path.join(TEMPLATES_DIR, 'SkilledWorker_MaintenanceFunds.docx')
    doc.save(filename)
    print(f'✅ Created: {filename}')
    return filename


if __name__ == '__main__':
    print('🚀 Generating Premium Templates...\n')
    
    create_spouse_proof_of_relationship()
    create_spouse_financial_evidence()
    create_skilled_worker_maintenance()
    
    print(f'\n✅ COMPLETE! Templates saved to: {TEMPLATES_DIR}')
    print('📝 Next steps: Review templates, then generate remaining 16 templates.')
